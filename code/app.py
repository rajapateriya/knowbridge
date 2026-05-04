"""
Web UI for KnowBridge — RAG-powered knowledge assistant.

Tab 1 — Knowledge Base: upload .md/.docx files, smart re-index (only changed/new files).
Tab 2 — Chat: ask questions against the indexed knowledge base.
"""

import os
import sys
import logging
import base64
import shutil
import subprocess
import tempfile

import gradio as gr
from dotenv import load_dotenv

from paths import APP_CONFIG_FPATH, PROMPT_CONFIG_FPATH, VECTOR_DB_DIR, LOGO_FPATH
from utils import load_yaml_config
from chat_history_db import ChatHistoryDB
from document_indexer import (
    upsert_document,
    list_indexed_sources,
)
from rag_pipeline import collection, respond_to_query, summarize_history, setup_logging

# ---------------------------------------------------------------------------
# Bootstrap
# ---------------------------------------------------------------------------

load_dotenv()
os.environ["TOKENIZERS_PARALLELISM"] = "false"

app_config = load_yaml_config(APP_CONFIG_FPATH)
prompt_config = load_yaml_config(PROMPT_CONFIG_FPATH)

# Configure logging based on config.yaml
setup_logging(log_level=app_config.get("log_level", "INFO"))
logging.info("Logging configured (console=%s)", app_config.get("log_level", "INFO"))

rag_prompt_config = prompt_config["rag_assistant_prompt"]
vectordb_params: dict = app_config["vectordb"]
llm_model: str = app_config["llm"]
_memory_cfg: dict = app_config.get("memory_strategies", {})
_trim_window: int = _memory_cfg.get("trimming_window_size", 6)
# summarize_after: total messages before rolling summarization triggers (default: 2x window)
_summarize_after: int = _memory_cfg.get("summarize_after", _trim_window * 2)

# Use the single collection instance from rag_pipeline.

# Shared SQLite chat history DB — one file, all sessions/users separated by session_id
db = ChatHistoryDB()

# ---------------------------------------------------------------------------
# Knowledge Base helpers
# ---------------------------------------------------------------------------

def _display_filename(source: str) -> str:
    """Convert a stored source key into a human-friendly filename."""
    # Back-compat: historically we stored markdown sources without extension.
    _, ext = os.path.splitext(source)
    if ext:
        return source
    return f"{source}.md"


def _extract_docx_text(file_path: str) -> str:
    """Extract plain text from a .docx file."""
    try:
        from docx import Document  # python-docx
    except Exception as e:  # pragma: no cover
        raise RuntimeError(
            "Missing dependency for .docx support. Install `python-docx` and retry."
        ) from e

    doc = Document(file_path)
    parts: list[str] = []

    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text:
            parts.append(text)

    # Include table content (common in policy docs)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = "\t".join([c for c in cells if c])
            if line:
                parts.append(line)

    return "\n".join(parts).strip()


def _extract_doc_text(file_path: str) -> str:
    """Extract plain text from a legacy .doc file.

    Strategy:
    - macOS: use the built-in `textutil` CLI when available.
    - Fallback: use LibreOffice (`soffice`) headless conversion to .txt.
    """
    # 1) macOS built-in converter (best UX: no Python deps)
    if sys.platform == "darwin" and shutil.which("textutil"):
        proc = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", file_path],
            capture_output=True,
            check=False,
        )
        if proc.returncode == 0:
            return proc.stdout.decode("utf-8", errors="replace").strip()

    # 2) LibreOffice conversion (cross-platform if installed)
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        with tempfile.TemporaryDirectory() as tmpdir:
            proc = subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--convert-to",
                    "txt:Text",
                    "--outdir",
                    tmpdir,
                    file_path,
                ],
                capture_output=True,
                check=False,
            )
            # LibreOffice typically writes <basename>.txt
            base = os.path.splitext(os.path.basename(file_path))[0]
            candidate = os.path.join(tmpdir, f"{base}.txt")
            if os.path.exists(candidate):
                with open(candidate, "r", encoding="utf-8", errors="replace") as f:
                    return f.read().strip()

            # Last resort: find any .txt output
            for name in os.listdir(tmpdir):
                if name.lower().endswith(".txt"):
                    with open(
                        os.path.join(tmpdir, name),
                        "r",
                        encoding="utf-8",
                        errors="replace",
                    ) as f:
                        return f.read().strip()

            stderr = (proc.stderr or b"").decode("utf-8", errors="replace").strip()
            raise RuntimeError(
                "LibreOffice conversion ran but produced no .txt output."
                + (f" Details: {stderr}" if stderr else "")
            )

    raise ValueError(
        "Unsupported file type: .doc (legacy Word). "
        "On macOS, install/enable `textutil` (usually built-in) or install LibreOffice (soffice). "
        "Alternatively, convert the file to .docx and re-upload."
    )


def _read_uploaded_text(file_path: str, fname: str) -> str:
    ext = os.path.splitext(fname)[1].lower()
    if ext == ".md":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    if ext == ".docx":
        return _extract_docx_text(file_path)
    if ext == ".doc":
        return _extract_doc_text(file_path)
    raise ValueError(f"Unsupported file type: {ext}")

def _sources_table() -> str:
    """Return a Markdown table of all indexed files."""
    sources = list_indexed_sources(collection)
    if not sources:
        return "_No files indexed yet._"
    header = "| File | Chunks | Hash (preview) |\n|------|--------|----------------|\n"
    rows = "".join(
        f"| {_display_filename(s['source'])} | {s['chunks']} | {s['file_hash'][:12]}... |\n"
        for s in sources
    )
    return header + rows


def upload_files(files) -> tuple[str, list[list]]:
    """
    Handle file uploads. For each file:
    - added     → new file, indexed fresh
    - updated   → file changed, old chunks removed and reindexed
    - unchanged → identical to what's already in the DB, skipped
    """
    if not files:
        return "No files selected.", _sources_table()

    results = []
    for file in files:
        try:
            # Gradio 6 passes FileData objects with .path (temp path) and .orig_name
            # Older Gradio passed objects where .name was the temp path
            if hasattr(file, "orig_name") and file.orig_name:
                fname = os.path.basename(file.orig_name)
            elif hasattr(file, "name"):
                fname = os.path.basename(file.name)
            else:
                fname = os.path.basename(str(file))

            file_path = file.path if hasattr(file, "path") else file.name

            ext = os.path.splitext(fname)[1].lower()
            if ext not in (".md", ".docx", ".doc"):
                results.append(
                    f"⚠ Skipped {fname} — only .md, .docx, and .doc files are supported."
                )
                continue

            # Keep backward-compatible IDs for markdown, but make .docx unique.
            source_name = (
                os.path.splitext(fname)[0] if ext == ".md" else fname
            )

            content = _read_uploaded_text(file_path, fname)
            if not content.strip():
                results.append(f"⚠ Skipped {fname} — no extractable text found.")
                continue

            status = upsert_document(collection, content, source_name)

            icon = {"added": "✅", "updated": "🔄", "unchanged": "⏭"}.get(status, "?")
            results.append(f"{icon} {fname} — {status}")

        except Exception as e:
            fname_display = fname if "fname" in dir() else str(file)
            results.append(f"❌ {fname_display} — error: {e}")

    summary = "\n".join(results)
    return summary, _sources_table()


def refresh_table() -> str:
    return _sources_table()


# ---------------------------------------------------------------------------
# Chat helpers
# ---------------------------------------------------------------------------

def _get_session_id(session_hash: str, username: str | None) -> str:
    """Return username if set, else fall back to the Gradio session hash."""
    name = (username or "").strip()
    return name if name else session_hash


def load_session(username: str, request: gr.Request) -> tuple[list[dict], str]:
    """
    Load persisted history from SQLite for the resolved session.
    Called on page load and when the user changes/submits their name.
    """
    session_id = _get_session_id(request.session_hash, username)
    history = db.load_history(session_id)
    display = [m for m in history if m["role"] in ("user", "assistant")]
    label = f"Session: **{session_id}**"
    return display, label


def chat(
    message: str,
    history: list[dict],
    username: str,
    assistant_type: str,
    request: gr.Request,
) -> tuple[str, list[dict]]:
    """Handle a single chat turn with DB persistence, trimming, and summarization."""
    if not message.strip():
        return "", history

    session_id = _get_session_id(request.session_hash, username)
    assistant_type = (assistant_type or "PMO").strip() or "PMO"

    # 1. Load full persisted history for context building
    all_history = db.load_history(session_id)

    # 2. Build LLM context: rolling summary + trimmed recent window
    summary = db.get_summary(session_id)
    trimmed = all_history[-_trim_window:] if len(all_history) > _trim_window else all_history

    # 3. RAG pipeline
    response, source_files = respond_to_query(
        prompt_config=rag_prompt_config,
        query=message,
        llm=llm_model,
        chat_history=trimmed,
        history_summary=summary,
        assistant_type=assistant_type,
        **vectordb_params,
    )

    if source_files:
        sources_md = "\n\n---\n\U0001f4c4 **Referenced documents:** " + ", ".join(
            f"`{_display_filename(s)}`" for s in source_files
        )
        full_response = response + sources_md
    else:
        full_response = response

    # 4. Persist new messages
    db.save_message(session_id, "user", message)
    db.save_message(session_id, "assistant", full_response)

    # 5. Trigger rolling summarization once history grows past summarize_after
    all_history = db.load_history(session_id)
    if len(all_history) > _summarize_after:
        to_summarize = all_history[:-_trim_window]
        existing_summary = db.get_summary(session_id)
        try:
            new_summary = summarize_history(to_summarize, llm_model, existing_summary)
            db.save_summary(session_id, new_summary)
        except Exception as e:
            logging.warning(f"Summarization failed (non-fatal): {e}")

    # 6. Update Gradio display
    history.append({"role": "user", "content": message})
    history.append({"role": "assistant", "content": full_response})
    return "", history


def clear_chat(username: str, request: gr.Request) -> tuple[list, str]:
    """Clear Gradio display AND the persisted DB history for this session."""
    session_id = _get_session_id(request.session_hash, username)
    db.clear_history(session_id)
    return [], ""


# ---------------------------------------------------------------------------
# Gradio UI
# ---------------------------------------------------------------------------

_logo_b64 = base64.b64encode(open(LOGO_FPATH, "rb").read()).decode()

with gr.Blocks(title="KnowBridge", theme=gr.themes.Soft()) as demo:
    gr.HTML(
        f"""
        <div style="display:flex;align-items:center;gap:14px;padding:8px 0 4px 0;">
            <img src="data:image/png;base64,{_logo_b64}"
                 style="width:86px;height:86px;object-fit:contain;flex-shrink:0;">
            <div>
                <h1 style="margin:0;font-size:1.8rem;line-height:1.2;">KnowBridge</h1>
                <p style="margin:2px 0 0 0;color:#555;">Transform your documents into an intelligent knowledge assistant.
                Upload, explore, and get instant answers powered by AI.</p>
            </div>
        </div>
        """
    )

    with gr.Tabs():

        # ── Tab 1: Knowledge Base ────────────────────────────────────────────
        with gr.Tab("📁 Knowledge Base"):
            gr.Markdown("### Upload Documents")
            gr.Markdown(
                "Upload `.md`, `.docx`, and `.doc` files. Files are automatically compared with what's already "
                "indexed — only **new** or **changed** files are re-embedded."
            )

            with gr.Row():
                file_input = gr.File(
                    label="Select .md, .docx, or .doc files",
                    file_count="multiple",
                    file_types=[".md", ".docx", ".doc"],
                )

            upload_btn = gr.Button("Upload & Index", variant="primary")
            upload_status = gr.Textbox(
                label="Upload status",
                lines=5,
                interactive=False,
            )

            gr.Markdown("### Indexed Files")
            indexed_table = gr.Markdown(value=_sources_table())
            refresh_btn = gr.Button("Refresh list")

            gr.Markdown("### Assistant Configuration")
            assistant_type_input = gr.Textbox(
                label="Assistant type",
                placeholder="e.g. PMO, HR, Finance, Legal...",
                value="PMO",
            )

            upload_btn.click(
                fn=upload_files,
                inputs=[file_input],
                outputs=[upload_status, indexed_table],
            )
            refresh_btn.click(
                fn=refresh_table,
                inputs=[],
                outputs=[indexed_table],
            )

        # ── Tab 2: Chat ──────────────────────────────────────────────────────
        with gr.Tab("💬 Chat"):
            gr.Markdown("### Ask a question")
            gr.Markdown(
                "Questions are answered using only the indexed documents. "
                "Enter your name to persist and resume conversations across sessions."
            )

            with gr.Row():
                username_input = gr.Textbox(
                    label="Your name / session (optional)",
                    placeholder="Enter a name to save and resume history. Leave blank for a temporary session.",
                )
            session_label = gr.Markdown("Session: **auto**")

            chatbot = gr.Chatbot(height=460, show_label=False)
            with gr.Row():
                msg_input = gr.Textbox(
                    placeholder="Type your question here...",
                    show_label=False,
                    scale=9,
                )
                send_btn = gr.Button("Send", variant="primary", scale=1)

            clear_btn = gr.Button("Clear chat")

            # Load history when user sets/changes their session name
            username_input.submit(
                fn=load_session,
                inputs=[username_input],
                outputs=[chatbot, session_label],
            )
            send_btn.click(
                fn=chat,
                inputs=[msg_input, chatbot, username_input, assistant_type_input],
                outputs=[msg_input, chatbot],
            )
            msg_input.submit(
                fn=chat,
                inputs=[msg_input, chatbot, username_input, assistant_type_input],
                outputs=[msg_input, chatbot],
            )
            clear_btn.click(
                fn=clear_chat,
                inputs=[username_input],
                outputs=[chatbot, msg_input],
            )

    # On every page load: refresh the indexed-files table AND restore the
    # chat history for the resolved session (uses gr.Request.session_hash).
    demo.load(fn=refresh_table, outputs=[indexed_table])
    demo.load(fn=load_session, inputs=[username_input], outputs=[chatbot, session_label])


if __name__ == "__main__":
    demo.launch(share=False)
